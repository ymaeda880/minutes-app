# lib/docx_minutes_export.py
# ------------------------------------------------------------
# 議事録テキスト（Markdown 風） → 美しい Word(.docx) 変換ヘルパー
# ------------------------------------------------------------
from __future__ import annotations

from io import BytesIO
from typing import Optional, Tuple, List

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
#  箇条書き用ヘルパー
# ============================================================
def _add_custom_bullet(doc: Document, text: str, level: int = 1, base_font: str = "游ゴシック") -> None:
    """
    Word の自動箇条書きに依存せず、
    行頭記号とインデントを完全制御して箇条書きを作る。

    level:
      1 -> ●（小さめ）
      2 -> ○（小さめ）
      3 -> ・
    """
    p = doc.add_paragraph(style="Normal")
    pf = p.paragraph_format

    if level <= 1:
        left = Pt(0)
        bullet_char = "●"
    elif level == 2:
        left = Pt(12)   # 約 4mm
        bullet_char = "○"
    else:
        left = Pt(24)
        bullet_char = "・"

    pf.left_indent = left
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.2

    # ● / ○ だけ少し小さめに描画して、その後ろに本文を通常サイズで入れる
    bullet_run = p.add_run(bullet_char + " ")
    bullet_run.font.name = base_font
    bullet_run.font.size = Pt(9)     # ← ここで記号だけ小さくする

    text_run = p.add_run(text)
    text_run.font.name = base_font
    text_run.font.size = Pt(11)      # 本文は通常サイズ


def _parse_bullet(line: str) -> Optional[Tuple[int, str]]:
    """
    行頭のスペース（半角・全角）と記号を見て
    「これは箇条書きか？」「レベルはいくつか？」を判定する。

    対応パターンの例:
      "- 本文"
      "  - 本文"
      "    - 本文"
      "　- 本文"   （全角スペース＋ハイフン）
      "・ 本文"
      "　・ 本文"
      "－ 本文"   （全角ハイフン）
    """
    i = 0
    half = 0
    full = 0
    while i < len(line) and line[i] in (" ", "\t", "　"):
        if line[i] == "　":
            full += 1
        else:
            half += 1
        i += 1

    rest = line[i:]

    # マーカー（半角/全角ハイフン、中黒）を許容
    if rest.startswith("- "):
        marker_len = 2
    elif rest.startswith("－ "):
        marker_len = 2
    elif rest.startswith("・ "):
        marker_len = 2
    else:
        return None

    # 全角スペースは半角2個ぶんでカウント
    indent_units = half + 2 * full

    if indent_units >= 4:
        level = 3
    elif indent_units >= 2:
        level = 2
    else:
        level = 1

    text = rest[marker_len:].strip()
    return level, text


def _add_horizontal_rule(doc: Document) -> None:
    """Markdown の --- を Word の段落下罫線に変換する。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)

    p_elm = p._p
    pPr = p_elm.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)


# ============================================================
#  Markdown 風テーブル検出・変換ヘルパー（GPT 出力前提）
# ============================================================

import re

def _is_separator_line(line: str) -> bool:
    """
    GPT が生成する Markdown 表の区切り行（| --- | --- | など）を
    できる限り広く許容して判定する。

    ※ ゆらぎの例:
        - |---|---|
        - | --- | --- |
        - ｜ --- ｜ --- ｜（全角パイプ混在）
        - | — | — |
        - | ーーー | ーーー |
        - （セル間の空白数が不定）
        - （ゼロ幅スペース混入）
    """

    # 全角パイプ → 半角に
    s = line.replace("｜", "|")

    # ゼロ幅スペースなどを除去
    s = s.replace("\u200b", "").replace("\u2060", "").replace("\ufeff", "")

    s = s.strip()

    # 最低限 | を2つ以上含まないと表ではない
    if s.count("|") < 2:
        return False

    # 中身をセルごとに見る
    cells = [c.strip() for c in s.strip("|").split("|")]

    if not cells:
        return False

    for c in cells:
        # 区切りセルは - : のみ、または全角/ダッシュ類のみを許容
        cc = c.replace(" ", "")
        # エムダッシュ・全角ダッシュなども - 扱いに正規化
        cc = cc.replace("—", "-").replace("–", "-").replace("ー", "-")

        if not cc:
            return False

        if not re.fullmatch(r"[-:]+", cc):
            return False

    return True


def _split_md_cells(line: str) -> List[str]:
    """
    1 行の Markdown 風テーブル行をセルに分割する。
    - 先頭/末尾の | は無視
    - セル内の \| は「|」として扱う
    """
    s = line.strip("\n").rstrip()
    # 先頭・末尾の | は削る（あれば）
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]

    cells: List[str] = []
    buf: List[str] = []
    escaped = False

    for ch in s:
        if escaped:
            # 直前がバックスラッシュならそのまま文字として追加
            buf.append(ch)
            escaped = False
            continue

        if ch == "\\":
            escaped = True
            continue

        if ch == "|":
            # セル区切り
            cell = "".join(buf).strip()
            cells.append(cell if cell != "" else " ")
            buf = []
        else:
            buf.append(ch)

    # 最後のセル
    cell = "".join(buf).strip()
    cells.append(cell if cell != "" else " ")

    return cells


def _is_gpt_table_block(lines: List[str]) -> bool:
    """
    GPT 出力を前提としたテーブルブロック判定。

    条件:
      - 2 行以上
      - 各行に半角 '|' が含まれる
      - 各行を _split_md_cells したときの列数がすべて同じ（2 列以上）
    """
    if len(lines) < 2:
        return False

    col_counts: List[int] = []

    for ln in lines:
        if "|" not in ln:
            return False  # パイプがない行が混ざっていたらテーブルではない

        cells = _split_md_cells(ln)
        if len(cells) < 2:
            return False  # 1 列だけならテーブルとみなさない
        col_counts.append(len(cells))

    # すべて同じ列数ならテーブルとみなす
    return len(set(col_counts)) == 1


def _add_table_from_lines(doc: Document, lines: List[str], base_font: str = "游ゴシック") -> None:
    """
    連続した Markdown 風テーブル行から Word の表を作成する。

    方針:
      - まずテーブルブロック全体が妥当かどうかを _is_gpt_table_block で確認
      - そのうえで「区切り行」（| --- | --- | など）はすべて捨てる
      - 残った行の先頭行をヘッダ、それ以降をデータ行として扱う
    """
    if not _is_gpt_table_block(lines):
        # フォールバック：すべて通常段落として出力
        for ln in lines:
            doc.add_paragraph(ln, style="Normal")
        return

    # --- 1) 区切り行を全部削除する ------------------------------
    content_lines: List[str] = [ln for ln in lines if not _is_separator_line(ln)]

    # すべて区切り行だった／おかしなケース → 段落として出す
    if not content_lines:
        for ln in lines:
            doc.add_paragraph(ln, style="Normal")
        return

    # --- 2) 先頭行をヘッダ、それ以外をデータ行として扱う ---------
    header_line = content_lines[0]
    data_lines = content_lines[1:]

    header_cells = _split_md_cells(header_line)
    ncols = len(header_cells)

    # 行数は ヘッダ 1 + データ行数（データ行が無くてもヘッダだけの1行表を作る）
    nrows = 1 + len(data_lines)
    if nrows <= 1:
        nrows = 1

    table = doc.add_table(rows=nrows, cols=ncols)
    try:
        table.style = "Table Grid"
    except Exception:
        # スタイルがなければ無視
        pass

    # --- ヘッダ行 ---
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        para = cell.paragraphs[0]
        para.text = ""  # 既存テキストをクリア
        run = para.add_run(cell_text)
        run.font.name = base_font
        run.font.size = Pt(10)
        run.bold = True

    # --- データ行 ---
    for i, line in enumerate(data_lines, start=1):
        row_cells = _split_md_cells(line)
        for j in range(ncols):
            text = row_cells[j] if j < len(row_cells) else " "
            cell = table.rows[i].cells[j]
            para = cell.paragraphs[0]
            para.text = ""
            run = para.add_run(text)
            run.font.name = base_font
            run.font.size = Pt(10)



# ============================================================
#  メイン関数：minutes 用 docx を組み立てる
# ============================================================
def build_minutes_docx(
    final_text: str,
    *,
    base_font: str = "游ゴシック",
) -> BytesIO:
    """
    議事録のプレーンテキスト（Markdown 風）から
    完全制御された見た目の .docx を生成して BytesIO で返す。

    - # 見出し  → Heading 1
    - ## 見出し → Heading 2
    - 箇条書き  → ● / ○ / ・ のカスタム箇条書き
    - ---       → 下線罫線
    - Markdown 風テーブルブロック → Word 表
    """
    doc = Document()

    # ---- スタイル定義（見た目をここで統一）----
    # Normal
    normal_style = doc.styles["Normal"]
    normal_style.font.name = base_font
    normal_style.font.size = Pt(11)
    normal_pf = normal_style.paragraph_format
    normal_pf.space_before = Pt(0)
    normal_pf.space_after = Pt(6)
    normal_pf.line_spacing = 1.3

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = base_font
    h1.font.bold = True
    h1.font.size = Pt(16)
    h1_pf = h1.paragraph_format
    h1_pf.space_before = Pt(18)
    h1_pf.space_after = Pt(6)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = base_font
    h2.font.bold = True
    h2.font.size = Pt(14)
    h2_pf = h2.paragraph_format
    h2_pf.space_before = Pt(12)
    h2_pf.space_after = Pt(4)

    prev_kind = "none"  # normal / bullet / h1 / h2 / hr / blank / table

    lines = final_text.splitlines()
    n = len(lines)
    i = 0

    while i < n:
        raw_line = lines[i]
        line = raw_line.rstrip()

        # 空行：直前が特定種別のときは抑制
        if line.strip() == "":
            if prev_kind in ("blank", "bullet", "h1", "h2", "hr", "table"):
                i += 1
                continue
            doc.add_paragraph("")
            prev_kind = "blank"
            i += 1
            continue

        # --- まずテーブルブロックを判定（最優先） ---
        if "|" in line:
            # 連続した「パイプを含む行」をまとめてテーブル候補ブロックとして取得
            block_lines: List[str] = []
            j = i
            while j < n:
                ln = lines[j]
                if ln.strip() == "":
                    break  # 空行でテーブル終端
                if "|" not in ln:
                    break  # パイプなし行でテーブル終端
                block_lines.append(ln)
                j += 1

            if _is_gpt_table_block(block_lines):
                _add_table_from_lines(doc, block_lines, base_font=base_font)
                prev_kind = "table"
                i = j
                continue
            # テーブルでない場合は、そのまま通常処理に落とす（1 行ずつ）

        # 横線
        if line.strip() in ("---", "―――", "ーーー"):
            _add_horizontal_rule(doc)
            prev_kind = "hr"
            i += 1
            continue

        # 見出し1
        if line.startswith("# "):
            text = line[2:].strip()
            doc.add_paragraph(text, style="Heading 1")
            prev_kind = "h1"
            i += 1
            continue

        # 見出し2
        if line.startswith("## "):
            text = line[3:].strip()
            doc.add_paragraph(text, style="Heading 2")
            prev_kind = "h2"
            i += 1
            continue

        # 箇条書き
        parsed = _parse_bullet(line)
        if parsed is not None:
            level, text = parsed
            _add_custom_bullet(doc, text, level=level, base_font=base_font)
            prev_kind = "bullet"
            i += 1
            continue

        # 通常段落
        p = doc.add_paragraph(line, style="Normal")
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.3
        prev_kind = "normal"
        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
