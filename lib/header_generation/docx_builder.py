# -*- coding: utf-8 -*-
# minutes_app/lib/header_generation/docx_builder.py
# ============================================================
# 議事録ヘッダー Word 作成
# ============================================================

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


def set_cell_border(cell) -> None:
    """Word 表セルに罫線を設定する。"""

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_margins(
    cell,
    *,
    top: int = 40,
    bottom: int = 40,
    left: int = 80,
    right: int = 80,
) -> None:
    """Word 表セルの内側余白を小さくして表の高さを詰める。"""

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for name, value in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)

        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    center: bool = False,
) -> None:
    """Word 表セルに文字列を設定する。"""

    cell.text = ""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if center
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)
    run.font.name = "游明朝"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
    run.font.size = Pt(8.5)
    run.bold = bold

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_cell_border(cell)
    set_cell_margins(
        cell,
        top=40,
        bottom=40,
        left=80,
        right=80,
    )


def set_table_column_widths(
    table,
    widths_mm: list[float],
) -> None:
    """Word 表の列幅を設定する。"""

    for row in table.rows:
        for idx, width in enumerate(widths_mm):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Mm(width)


def guess_widths_mm(col_count: int) -> list[float]:
    """列数から Word 表の列幅を推定する。"""

    if col_count == 4:
        return [28, 86, 28, 24]

    if col_count == 3:
        return [70, 62, 34]

    if col_count == 2:
        return [80, 86]

    if col_count <= 0:
        return []

    return [166 / col_count] * col_count


def build_docx_bytes(
    blocks: list[dict[str, Any]],
) -> bytes:
    """表ブロックから議事録ヘッダー Word の bytes を作成する。"""

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
    style.font.size = Pt(10)

    for block_index, block in enumerate(blocks):
        title = str(block.get("title", "")).strip()
        header = list(block.get("header", []))
        data = list(block.get("data", []))

        if not header:
            continue

        if block_index > 0:
            doc.add_paragraph("")

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(2)

        r_title = p_title.add_run(title)
        r_title.font.name = "游明朝"
        r_title._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
        r_title.font.size = Pt(10)
        r_title.bold = False

        rows_count = 1 + len(data)
        cols_count = len(header)

        table = doc.add_table(
            rows=rows_count,
            cols=cols_count,
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        widths_mm = guess_widths_mm(cols_count)
        set_table_column_widths(table, widths_mm)

        for col_idx, text in enumerate(header):
            set_cell_text(
                table.cell(0, col_idx),
                str(text),
                bold=False,
                center=True,
            )

        for row_idx, row_values in enumerate(data, start=1):
            for col_idx in range(cols_count):
                value = (
                    row_values[col_idx]
                    if col_idx < len(row_values)
                    else ""
                )
                set_cell_text(
                    table.cell(row_idx, col_idx),
                    str(value),
                    bold=False,
                    center=False,
                )

    bio = io.BytesIO()
    doc.save(bio)

    return bio.getvalue()