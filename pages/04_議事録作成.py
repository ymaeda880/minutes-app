# ------------------------------------------------------------
# 📝 議事録作成（整形済みテキスト → 議事録）— modern専用・リトライなし版
# ------------------------------------------------------------
# 04_議事録作成.py
from __future__ import annotations

import time
from typing import Dict, Any
from io import BytesIO

import streamlit as st
import pandas as pd
from openai import OpenAI
from datetime import datetime

# ==== .docx 読み取り／書き出し（python-docx） ====
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# ==== 共通ユーティリティ ====
from lib.prompts import (
    MINUTES_MAKER,
    MINUTES_MANDATORY_MODES,  # ← 議事録の種類（逐語録/簡易/詳細など）
    MINUTES_STYLE,            # ← 見た目スタイル用グループを追加
    get_group,
    build_prompt,
)
from lib.tokens import extract_tokens_from_response, debug_usage_snapshot  # modern専用
from lib.costs import estimate_chat_cost_usd
from config.config import DEFAULT_USDJPY

from lib.explanation import render_minutes_maker_expander

# ========================== 共通設定 ==========================
st.set_page_config(page_title="④ 議事録作成", page_icon="📝", layout="wide")
st.title("議事録作成 — 逐語録から正式議事録へ")
render_minutes_maker_expander()

OPENAI_API_KEY = st.secrets.get("openai", {}).get("api_key") or st.secrets.get("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("OpenAI API Key が見つかりません。.streamlit/secrets.toml を確認してください。")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)

# ---- セッション初期化（表示が消えない用の保険）----
st.session_state.setdefault("minutes_final_output", "")

# ========================== モデル設定補助 ==========================
def supports_temperature(model_name: str) -> bool:
    """GPT-5系は temperature 変更不可（=1固定）。"""
    return not model_name.startswith("gpt-5")

# ========================== レイアウト（土台） ==========================
left, right = st.columns([1, 1], gap="large")

# ========================== 左カラム：プロンプト設定 ==========================
with left:
    st.subheader("プロンプト")

    # 内容系プロンプト（逐語録 / 簡易 / 詳細）
    group = get_group(MINUTES_MAKER)
    # 見た目スタイルプロンプト（ベーシック / 横線＋装飾 / Word見出し向け）
    style_group = get_group(MINUTES_STYLE)

    # --- mandatory モードの初期化 ---
    mode_options = list(MINUTES_MANDATORY_MODES.keys())
    if "minutes_mode" not in st.session_state:
        # デフォルトは「簡易議事録の作成」
        st.session_state["minutes_mode"] = "簡易議事録の作成"

    if "minutes_mandatory" not in st.session_state:
        st.session_state["minutes_mandatory"] = MINUTES_MANDATORY_MODES[
            st.session_state["minutes_mode"]
        ]

    def _on_change_minutes_mode() -> None:
        mode = st.session_state["minutes_mode"]
        st.session_state["minutes_mandatory"] = MINUTES_MANDATORY_MODES.get(
            mode,
            MINUTES_MANDATORY_MODES["簡易議事録の作成"],
        )

    # --- 議事録の種類（内容） ---
    st.radio(
        "議事録の種類",
        options=mode_options,
        key="minutes_mode",          # index は渡さない
        on_change=_on_change_minutes_mode,
        help="逐語録 / 逐語録のまとめ / 簡易議事録 / 詳細議事録 を切り替えます。",
    )

    # --- 見た目スタイルの選択（ベーシック / 横線 / Word見出し） ---
    style_labels = style_group.preset_labels()
    if "minutes_style_label" not in st.session_state:
        st.session_state["minutes_style_label"] = style_labels[2]
        

    st.radio(
        "見た目のスタイル",
        options=style_labels,
        key="minutes_style_label",
        help="Word に貼り付けたときのレイアウトや見た目を切り替えます。",
    )

    # --- プリセットなどの初期化（内容側） ---
    if "minutes_preset_label" not in st.session_state:
        st.session_state["minutes_preset_label"] = group.label_for_key(
            group.default_preset_key
        )
    if "minutes_preset_text" not in st.session_state:
        st.session_state["minutes_preset_text"] = group.body_for_label(
            st.session_state["minutes_preset_label"]
        )
    if "minutes_extra_text" not in st.session_state:
        st.session_state["minutes_extra_text"] = ""

    with st.expander("必須パート（編集可）", expanded=False):
        st.text_area(
            "必ず入る部分（常にプロンプトの先頭に含まれます）",
            height=220,
            key="minutes_mandatory",
        )

    # --- 追記プリセット（選択 → 本文を自動反映） ---
    def _on_change_preset():
        st.session_state["minutes_preset_text"] = group.body_for_label(
            st.session_state["minutes_preset_label"]
        )

    st.selectbox(
        "追記プリセット（内容）",
        options=group.preset_labels(),
        key="minutes_preset_label",   # index を使わない
        help="選んだ内容が上の必須文の下に自動的に連結されます。",
        on_change=_on_change_preset,
    )

    # 選択中のプリセット本文（編集可）
    st.text_area("（編集可）プリセット本文（内容）", height=120, key="minutes_preset_text")

    # 任意の追加指示
    st.text_area("追加指示（任意）", height=88, key="minutes_extra_text")

# ========================== 右カラム：入力テキスト ==========================
with right:
    st.subheader("整形済みテキスト（入力）")

    up = st.file_uploader(
        "③ページの整形結果（.txt または .docx）をアップロードするか、下の欄に貼り付けてください。",
        type=["txt", "docx"],
        accept_multiple_files=False,
    )

    if up is not None:
        if up.name.lower().endswith(".docx"):
            if not HAS_DOCX:
                st.error("`.docx` を読み込むには python-docx が必要です。`pip install python-docx` を実行してください。")
            else:
                data = up.read()
                try:
                    doc = Document(BytesIO(data))
                    text_from_file = "\n".join(p.text for p in doc.paragraphs)
                except Exception as e:
                    st.error(f"Wordファイルの読み込みに失敗しました: {e}")
                    text_from_file = ""
                st.session_state["minutes_source_text"] = text_from_file
        else:
            raw = up.read()
            try:
                text_from_file = raw.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text_from_file = raw.decode("cp932")
                except Exception:
                    text_from_file = raw.decode(errors="ignore")
            st.session_state["minutes_source_text"] = text_from_file

    src = st.text_area(
        "③ページの整形結果を引き継ぐか、ここに貼り付けてください。",
        value=st.session_state.get("minutes_source_text", ""),
        height=460,
        placeholder="「③ 話者分離・整形（新）」の結果を流し込む想定です。",
    )

# ========================== サイドバー：モデル設定＋通貨 ==========================
with st.sidebar:
    st.subheader("モデル設定")

    model = st.selectbox(
        "モデル",
        [
            "gpt-5",
            "gpt-5-mini",
            "gpt-5-nano",
            "gpt-4.1-mini",
            "gpt-4.1",
        ],
        index=1,
    )

    temp_supported = supports_temperature(model)
    temperature = st.slider(
        "温度（0=厳格 / 2=自由）",
        0.0,
        2.0,
        value=0.7,
        step=0.1,
        disabled=not temp_supported,
        help="GPT-5 系列は temperature=1 固定です",
    )
    if not temp_supported:
        st.caption("ℹ️ GPT-5 系列は temperature を変更できません（=1固定）")

    max_completion_tokens = st.slider(
        "最大出力トークン（目安）",
        min_value=1000,
        max_value=120000,
        value=100000,
        step=500,
        help="長めの議事録生成なら 8,000〜12,000 程度を推奨（本版はリトライなし）。",
    )

    st.subheader("通貨換算（任意）")
    usd_jpy = st.number_input(
        "USD/JPY",
        min_value=50.0,
        max_value=500.0,
        value=float(DEFAULT_USDJPY),
        step=0.5,
    )

# 実行ボタン（メイン側）
run_btn = st.button("📝 議事録を生成", type="primary", use_container_width=True)

# ========================== 実行（モデル呼び出し：リトライなし） ==========================
if run_btn:
    if not src.strip():
        st.warning("整形済みテキストを入力してください。")
    else:
        # --- 見た目スタイルの本文を取得 ---
        style_body = style_group.body_for_label(
            st.session_state.get(
                "minutes_style_label",
                style_group.label_for_key(style_group.default_preset_key),
            )
        )

        # --- 内容プリセット + 見た目スタイル を合体 ---
        base_preset = st.session_state.get("minutes_preset_text", "") or ""
        if style_body:
            merged_preset = (
                base_preset.strip()
                + "\n\n【見た目のスタイル指示】\n"
                + style_body.strip()
            )
        else:
            merged_preset = base_preset

        # --- プロンプト組み立て ---
        combined = build_prompt(
            st.session_state["minutes_mandatory"],   # 議事録の種類（逐語録/簡易/詳細）の必須部分
            merged_preset,                           # 内容プリセット + 見た目スタイル
            st.session_state["minutes_extra_text"],  # 任意の追加指示
            src,
        )



        def call_once(prompt_text: str, out_tokens: int):
            chat_kwargs: Dict[str, Any] = dict(
                model=model,
                messages=[{"role": "user", "content": prompt_text}],
                max_completion_tokens=int(out_tokens),
            )
            if temp_supported and abs(temperature - 1.0) > 1e-9:
                chat_kwargs["temperature"] = float(temperature)
            return client.chat.completions.create(**chat_kwargs)

        t0 = time.perf_counter()
        with st.spinner("議事録を生成中…"):
            resp = call_once(combined, max_completion_tokens)

            text = ""
            finish_reason = None
            if resp and getattr(resp, "choices", None):
                try:
                    text = resp.choices[0].message.content or ""
                except Exception:
                    text = getattr(resp.choices[0], "text", "")
                try:
                    finish_reason = resp.choices[0].finish_reason
                except Exception:
                    finish_reason = None

        elapsed = time.perf_counter() - t0

        if text.strip():
            st.session_state["minutes_final_output"] = text
            if finish_reason == "length":
                st.info(
                    "finish_reason=length: 出力が上限で切れています。必要に応じて最大出力トークンを増やしてください。"
                )
        else:
            st.warning("⚠️ モデルから空の応答が返されました。レスポンス全体を表示します。")
            try:
                st.json(resp.model_dump())
            except Exception:
                st.write(resp)

        if "resp" in locals():
            input_tok, output_tok, total_tok = extract_tokens_from_response(resp)
            usd = estimate_chat_cost_usd(model, input_tok, output_tok)
            jpy = (usd * usd_jpy) if usd is not None else None

            metrics_data = {
                "処理時間": [f"{elapsed:.2f} 秒"],
                "入力トークン": [f"{input_tok:,}"],
                "出力トークン": [f"{output_tok:,}"],
                "合計トークン": [f"{total_tok:,}"],
                "概算 (USD/JPY)": [
                    f"${usd:,.6f} / ¥{jpy:,.2f}" if usd is not None else "—"
                ],
            }
            st.subheader("トークンと料金の概要")
            st.table(pd.DataFrame(metrics_data))

            with st.expander("🔍 トークン算出の内訳（modern usage スナップショット）"):
                try:
                    st.write(debug_usage_snapshot(getattr(resp, "usage", None)))
                except Exception as e:
                    st.write({"error": str(e)})

# ========================== 生成結果の表示 ＆ ダウンロード ==========================
final_text = (st.session_state.get("minutes_final_output") or "").strip()

def safe_filename(s: str) -> str:
    bad = '\\/:*?"<>|'
    for ch in bad:
        s = s.replace(ch, "_")
    return s

if final_text:
    st.markdown("### 📝 生成結果（Markdown 表示）")
    st.markdown(final_text)

    st.subheader("📥 議事録の保存")

    # --- TXT 保存 ---
    txt_bytes = final_text.encode("utf-8")
    st.download_button(
        label="💾 テキストで保存 (.txt)",
        data=txt_bytes,
        file_name="minutes_output.txt",
        mime="text/plain",
        use_container_width=True,
        key="dl_txt_minutes",
    )


    # --- DOCX 保存（Markdown構造を Word スタイルに変換して綺麗に出力） ---
    if HAS_DOCX:
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            # ① モード（逐語録 / 簡易議事録 / 詳細議事録）
            mode_label = st.session_state.get("minutes_mode", "議事録")
            safe_label = safe_filename(mode_label)

            # ② 日時：YYYYMMDD_HHMM
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")

            # ③ Word 文書作成
            doc = Document()

            # ---- スタイル定義 ----
            # 通常段落
            normal_style = doc.styles["Normal"]
            normal_style.font.name = "游ゴシック"
            normal_style.font.size = Pt(11)

            # 見出し1
            h1 = doc.styles["Heading 1"]
            h1.font.name = "游ゴシック"
            h1.font.bold = True
            h1.font.size = Pt(14)

            # 見出し2
            h2 = doc.styles["Heading 2"]
            h2.font.name = "游ゴシック"
            h2.font.bold = True
            h2.font.size = Pt(12)

            # ---- 行単位で処理 ----
            for raw_line in final_text.splitlines():
                line = raw_line.rstrip()

                # 空行は空段落
                if line.strip() == "":
                    doc.add_paragraph("")
                    continue

                # 横線（---）を段落下罫線に変換
                if line.strip() in ("---", "―――", "ーーー"):
                    p = doc.add_paragraph()
                    pf = p.paragraph_format
                    pf.space_before = Pt(6)
                    pf.space_after = Pt(0)

                    # <w:pBdr> を自前で作る
                    p_elm = p._p  # paragraph の XML 要素
                    pPr = p_elm.get_or_add_pPr()
                    pBdr = pPr.find(qn("w:pBdr"))
                    if pBdr is None:
                        pBdr = OxmlElement("w:pBdr")
                        pPr.append(pBdr)

                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "6")      # 線の太さ
                    bottom.set(qn("w:space"), "1")   # 文字との間隔
                    bottom.set(qn("w:color"), "auto")
                    pBdr.append(bottom)
                    continue

                # 見出し1 (# 見出し)
                if line.startswith("# "):
                    text = line[2:].strip()
                    doc.add_paragraph(text, style="Heading 1")
                    continue

                # 見出し2 (## 見出し)
                if line.startswith("## "):
                    text = line[3:].strip()
                    doc.add_paragraph(text, style="Heading 2")
                    continue

                # 箇条書き (- 文)
                if line.startswith("- "):
                    text = line[2:].strip()
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(text)
                    # 少しだけ行間を詰めるなど調整したい場合
                    p.paragraph_format.space_after = Pt(0)
                    continue

                # 通常段落
                p = doc.add_paragraph(line, style="Normal")
                pf = p.paragraph_format
                pf.space_after = Pt(6)
                pf.line_spacing = 1.2

            # ---- 書き出し ----
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            # ---- ダウンロードボタン ----
            st.download_button(
                label="💾 Wordで保存 (.docx)",
                data=docx_buffer,
                file_name=f"minutes_{safe_label}_{timestamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_docx_minutes",
            )

        except Exception as e:
            st.error(f"Word 出力でエラーが発生しました: {e}")

